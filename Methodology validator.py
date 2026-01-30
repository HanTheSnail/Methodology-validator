import streamlit as st
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
from collections import defaultdict

# Question type requirements mapping
QTYPE_REQUIREMENTS = {
    'Single': {
        'options_required': True,
        'details_fields': {
            'required': ['kind', 'select_count'],
            'optional': ['none_text', 'other_text', 'randomise', 'goto', 'description']
        },
        'details_template': 'kind: exactly\nselect_count: 1\n',
        'description': 'Single choice question - exactly one option must be selected'
    },
    'Multi': {
        'options_required': True,
        'details_fields': {
            'required': ['kind', 'select_count'],
            'optional': ['randomise', 'none_text', 'other_text', 'goto', 'description']
        },
        'details_template': 'kind: atleast\nselect_count: 1\nrandomise: true\n',
        'description': 'Multiple choice - select at least N options'
    },
    'Grid': {
        'options_required': True,  # CRITICAL: Grid MUST have options or upload will fail
        'details_fields': {
            'required': [],
            'optional': ['comment_title', 'comment_description', 'comment_mandatory', 'randomise', 'goto', 'kind']
        },
        'details_template': 'comment_title: Would you like to add a comment?\ncomment_description: Tell us why\n',
        'description': 'Grid/Image grid - display options in grid layout. MUST have Options or upload fails!'
    },
    'Ranking': {
        'options_required': True,
        'details_fields': {
            'required': [],
            'optional': ['na_title', 'randomise', 'goto']
        },
        'details_template': 'na_title: Unranked\n',
        'description': 'Ranking question - order items by preference'
    },
    'Open': {
        'options_required': False,
        'details_fields': {
            'required': [],
            'optional': ['hint', 'min_length', 'goto', 'max_length']
        },
        'details_template': 'hint: Please explain\nmin_length: 0\n',
        'description': 'Open text response'
    },
    'Upload': {
        'options_required': False,
        'details_fields': {
            'required': ['type'],
            'optional': ['goto', 'description']
        },
        'details_template': 'type: photo\n',
        'description': 'File upload (photo/video/audio) - MUST specify type'
    },
    'Image Highlight': {
        'options_required': False,
        'details_fields': {
            'required': ['description'],
            'optional': ['hint', 'dynamic_watermark', 'goto']
        },
        'details_template': 'description: Click areas on the image\nhint: Tap to highlight\n',
        'description': 'Click/tap areas on image'
    },
    'Info Image': {
        'options_required': True,
        'details_fields': {
            'required': [],
            'optional': ['duration', 'dynamic_watermark']
        },
        'details_template': '',
        'description': 'Display information/images (Options: image descriptions with optional {text:} {url:})'
    },
    'AB': {
        'options_required': True,  # CRITICAL: AB MUST have options (typically 2) or upload will fail
        'details_fields': {
            'required': [],
            'optional': ['description', 'duration', 'goto']
        },
        'details_template': 'description: Choose A or B\n',
        'description': 'A/B comparison - MUST have exactly 2 options or upload fails!'
    },
    'Swipe': {
        'options_required': True,
        'details_fields': {
            'required': [],
            'optional': ['description', 'overlay', 'goto']
        },
        'details_template': 'description: Swipe direction\noverlay: Swipe text\n',
        'description': 'Swipe gesture (typically Left/Right)'
    },
    'Group': {
        'options_required': False,
        'details_fields': {
            'required': [],
            'optional': ['randomize']
        },
        'details_template': 'randomize: true\n',
        'description': 'Question group container - questions inside numbered as X.1, X.2, etc.'
    },
    'Title': {
        'options_required': False,
        'details_fields': {'required': [], 'optional': []},
        'details_template': '',
        'description': 'Survey title'
    },
    'Description': {
        'options_required': False,
        'details_fields': {'required': [], 'optional': []},
        'details_template': '',
        'description': 'Survey description/brief'
    },
    'Info': {
        'options_required': False,
        'details_fields': {'required': [], 'optional': ['duration']},
        'details_template': '',
        'description': 'Information display (non-interactive)'
    }
}

def parse_details(details_text):
    """Parse Details column into structured format"""
    if not details_text or not details_text.strip():
        return {}
    
    details_dict = {}
    lines = details_text.strip().split('\n')
    
    for line in lines:
        if ':' in line:
            key, value = line.split(':', 1)
            details_dict[key.strip()] = value.strip()
    
    return details_dict

def infer_question_type(qtype, content, options, details):
    """Attempt to infer the question type from content, options, and details"""
    if qtype and qtype.strip() and qtype in QTYPE_REQUIREMENTS:
        return qtype, None
    
    details_dict = parse_details(details)
    inferred = None
    confidence = ""
    
    # Strong indicators
    if 'type' in details_dict and details_dict['type'] in ['photo', 'video', 'audio']:
        return 'Upload', 'HIGH (has type: photo/video/audio)'
    
    # Grid can have comment fields but MUST have options (text or will be images)
    if 'comment_title' in details_dict or 'comment_description' in details_dict:
        if options and options.strip():
            return 'Grid', 'HIGH (has comment fields + options)'
        else:
            # Has grid indicators but no options - this will FAIL upload
            return 'Grid', 'HIGH (has comment fields) - WARNING: NEEDS OPTIONS OR UPLOAD WILL FAIL'
    
    # Single vs Multi based on kind and select_count
    if 'kind' in details_dict and 'select_count' in details_dict:
        if options and options.strip():
            if details_dict.get('kind') == 'exactly' and details_dict.get('select_count') == '1':
                return 'Single', 'HIGH (kind: exactly, select_count: 1)'
            else:
                return 'Multi', 'MEDIUM (has kind/select_count but not Single pattern)'
    
    # Check for specific keywords in content
    content_lower = content.lower() if content else ''
    
    if 'upload' in content_lower and ('photo' in content_lower or 'video' in content_lower or 'image' in content_lower):
        return 'Upload', 'MEDIUM (content mentions upload)'
    
    if 'rank' in content_lower or 'order' in content_lower:
        if options and options.strip():
            return 'Ranking', 'MEDIUM (content mentions ranking/ordering)'
    
    # Check if it's an info/display question (no interaction expected)
    if not details or not details.strip():
        if not options or not options.strip():
            if content and len(content) > 50:
                return 'Info', 'LOW (long content, no options/details - might be info display)'
    
    # If has options but no details indicators
    if options and options.strip():
        option_count = len([o for o in options.split('\n') if o.strip()])
        
        if option_count <= 2:
            return 'AB', 'LOW (2 options, might be A/B choice)'
        else:
            # Default to Single if we have options
            return 'Single', 'LOW (has options, defaulting to Single)'
    
    # If no options and asks a question
    if content and ('?' in content or 'why' in content_lower or 'what' in content_lower or 'how' in content_lower):
        return 'Open', 'MEDIUM (question with no options)'
    
    return None, None

def validate_question(qtype, qid, content, options, details):
    """Validate a question row and return issues"""
    issues = []
    warnings = []
    suggested_type = None
    confidence = None
    original_qtype = qtype
    
    # Skip header row
    if qtype in ['Type'] or qid in ['QID', '']:
        return issues, warnings, {}, None, None
    
    # Try to infer type if missing
    if not qtype or not qtype.strip():
        if qid and qid.strip():
            suggested_type, confidence = infer_question_type(qtype, content, options, details)
            if suggested_type:
                issues.append(f"❌ Type is MISSING")
                # Use suggested type for further validation
                qtype = suggested_type
            else:
                issues.append("❌ Type is MISSING - Cannot infer type from content")
                return issues, warnings, {}, suggested_type, confidence
        else:
            return issues, warnings, {}, None, None
    
    # Get requirements for this question type
    if qtype not in QTYPE_REQUIREMENTS:
        # Check if it looks like a typo or unknown type
        if qtype and qtype.strip():
            issues.append(f"❌ Unknown question type: '{qtype}'")
        return issues, warnings, {}, suggested_type, confidence
    
    reqs = QTYPE_REQUIREMENTS[qtype]
    
    # Check if Options is required but missing
    if reqs['options_required']:
        if not options or not options.strip():
            issues.append(f"❌ Options are REQUIRED for {qtype} questions but are missing")
    
    # Parse existing details
    details_dict = parse_details(details)
    
    # Check required Detail fields
    missing_required = []
    for req_field in reqs['details_fields']['required']:
        if req_field not in details_dict:
            missing_required.append(req_field)
    
    if missing_required:
        issues.append(f"❌ Required Details fields missing: {', '.join(missing_required)}")
    
    # Check for common optional fields that might be missing
    suggested_optional = []
    for opt_field in reqs['details_fields']['optional']:
        if opt_field not in details_dict:
            # Highlight particularly common ones
            if opt_field in ['description', 'hint'] and qtype in ['Single', 'Multi', 'Grid', 'Open', 'Image Highlight']:
                warnings.append(f"⚠️ Consider adding '{opt_field}' in Details")
    
    # Special validations
    if qtype in ['Single', 'Multi']:
        if 'kind' in details_dict:
            valid_kinds = ['exactly', 'atleast', 'atmost']
            if details_dict['kind'] not in valid_kinds:
                issues.append(f"❌ 'kind' must be one of: {', '.join(valid_kinds)}")
        
        if 'select_count' in details_dict:
            try:
                count = int(details_dict['select_count'])
                if count < 1:
                    issues.append("❌ 'select_count' must be at least 1")
            except ValueError:
                issues.append("❌ 'select_count' must be a number")
    
    if qtype == 'Upload':
        if 'type' in details_dict:
            valid_types = ['photo', 'video', 'audio']
            if details_dict['type'] not in valid_types:
                issues.append(f"❌ Upload 'type' must be one of: {', '.join(valid_types)}")
    
    if qtype == 'AB' and options:
        option_count = len([o for o in options.split('\n') if o.strip()])
        if option_count != 2:
            warnings.append(f"⚠️ AB questions typically have exactly 2 options (found {option_count})")
    
    return issues, warnings, reqs, suggested_type, confidence

def analyze_document(doc):
    """Analyze the document and return validation results"""
    results = []
    table = None
    
    # Find the first table with correct headers
    for t in doc.tables:
        if len(t.rows) > 0:
            headers = [cell.text.strip() for cell in t.rows[0].cells]
            if 'Type' in headers and 'QID' in headers:
                table = t
                break
    
    if not table:
        return None, "No valid methodology table found (must have Type and QID columns)"
    
    # Validate each row
    for idx, row in enumerate(table.rows[1:], start=1):  # Skip header
        cells = [cell.text.strip() for cell in row.cells]
        
        if len(cells) >= 5:
            qtype, qid, content, options, details = cells[0], cells[1], cells[2], cells[3], cells[4]
            
            # Skip completely empty rows
            if not any([qtype, qid, content, options, details]):
                continue
            
            issues, warnings, reqs, suggested_type, confidence = validate_question(qtype, qid, content, options, details)
            
            if issues or warnings:
                results.append({
                    'row': idx,
                    'qid': qid,
                    'qtype': qtype if qtype else suggested_type,
                    'suggested_type': suggested_type,
                    'confidence': confidence,
                    'issues': issues,
                    'warnings': warnings,
                    'requirements': reqs
                })
    
    return results, None

def add_suggestions_to_document(doc, validation_results):
    """Add a suggestions table after the main table"""
    
    # Add a page break
    doc.add_page_break()
    
    # Add heading
    heading = doc.add_paragraph()
    heading_run = heading.add_run("📋 VALIDATION REPORT & SUGGESTIONS")
    heading_run.bold = True
    heading_run.font.size = Pt(16)
    heading_run.font.color.rgb = RGBColor(0, 51, 153)
    
    doc.add_paragraph()
    
    # Add summary
    total_issues = sum(len(r['issues']) for r in validation_results)
    total_warnings = sum(len(r['warnings']) for r in validation_results)
    
    summary = doc.add_paragraph()
    summary.add_run(f"Found {total_issues} issues and {total_warnings} warnings across {len(validation_results)} questions.\n\n")
    
    # Create detailed report table
    for result in validation_results:
        # Question header
        q_para = doc.add_paragraph()
        display_type = result['qtype'] if result['qtype'] else "MISSING"
        q_run = q_para.add_run(f"Question {result['qid']} (Row {result['row']}) - Type: {display_type}")
        q_run.bold = True
        q_run.font.size = Pt(12)
        
        # Show suggested type if available
        if result.get('suggested_type'):
            suggest_para = doc.add_paragraph()
            suggest_run = suggest_para.add_run(f"💡 Suggested Type: {result['suggested_type']}")
            suggest_run.font.color.rgb = RGBColor(0, 153, 0)
            suggest_run.bold = True
            if result.get('confidence'):
                suggest_para.add_run(f" (Confidence: {result['confidence']})")
            suggest_para.paragraph_format.left_indent = Pt(20)
        
        if result['qtype'] and result['qtype'] in QTYPE_REQUIREMENTS:
            desc_para = doc.add_paragraph()
            desc_para.add_run(f"ℹ️ {QTYPE_REQUIREMENTS[result['qtype']]['description']}")
            desc_para.paragraph_format.left_indent = Pt(20)
        
        # Issues
        if result['issues']:
            for issue in result['issues']:
                issue_para = doc.add_paragraph(issue, style='List Bullet')
                issue_para.paragraph_format.left_indent = Pt(20)
                for run in issue_para.runs:
                    run.font.color.rgb = RGBColor(204, 0, 0)
        
        # Warnings
        if result['warnings']:
            for warning in result['warnings']:
                warn_para = doc.add_paragraph(warning, style='List Bullet')
                warn_para.paragraph_format.left_indent = Pt(20)
                for run in warn_para.runs:
                    run.font.color.rgb = RGBColor(255, 153, 0)
        
        # Suggested template
        if result['requirements'] and result['requirements'].get('details_template'):
            template_para = doc.add_paragraph()
            template_para.add_run("💡 Suggested Details template:\n").bold = True
            template_para.paragraph_format.left_indent = Pt(20)
            
            code_para = doc.add_paragraph(result['requirements']['details_template'])
            code_para.paragraph_format.left_indent = Pt(40)
            for run in code_para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 102, 0)
        
        # Required and optional fields
        if result['requirements']:
            fields_para = doc.add_paragraph()
            fields_para.paragraph_format.left_indent = Pt(20)
            
            if result['requirements']['details_fields']['required']:
                fields_para.add_run("Required: ").bold = True
                fields_para.add_run(", ".join(result['requirements']['details_fields']['required']))
                fields_para.add_run("\n")
            
            if result['requirements']['details_fields']['optional']:
                fields_para.add_run("Optional: ").bold = True
                fields_para.add_run(", ".join(result['requirements']['details_fields']['optional']))
        
        doc.add_paragraph()  # Spacing
    
    return doc

st.set_page_config(page_title="Questionnaire Methodology Validator", page_icon="📋", layout="wide")

st.title("📋 Questionnaire Methodology Validator")
st.markdown("""
Upload your methodology Word document and this tool will:
- ✅ Identify missing Type values
- ✅ Validate required Details fields for each question type
- ✅ Check Options requirements
- ✅ Suggest templates for Details based on question type
- ✅ Preserve ALL original formatting

**The tool generates a report showing exactly what's missing without modifying your original table.**
""")

# Question type reference
with st.expander("📖 Question Type Reference Guide"):
    for qtype, info in QTYPE_REQUIREMENTS.items():
        if qtype not in ['Title', 'Description']:
            st.markdown(f"**{qtype}**")
            st.markdown(f"_{info['description']}_")
            
            if info['details_fields']['required']:
                st.markdown(f"🔴 Required in Details: `{', '.join(info['details_fields']['required'])}`")
            if info['details_fields']['optional']:
                st.markdown(f"🟡 Optional in Details: `{', '.join(info['details_fields']['optional'])}`")
            
            if info['details_template']:
                st.code(info['details_template'], language=None)
            
            st.markdown("---")

uploaded_file = st.file_uploader("Upload Methodology Document (.docx)", type=['docx'])

if uploaded_file:
    # Load document
    doc = Document(uploaded_file)
    
    st.success("✅ Document loaded successfully!")
    
    # Analyze
    with st.spinner("Analyzing document..."):
        validation_results, error = analyze_document(doc)
    
    if error:
        st.error(error)
    elif not validation_results:
        st.success("🎉 No issues found! Your methodology looks good.")
    else:
        # Display summary
        total_issues = sum(len(r['issues']) for r in validation_results)
        total_warnings = sum(len(r['warnings']) for r in validation_results)
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Questions with Issues", len(validation_results))
        with col2:
            st.metric("Critical Issues", total_issues)
        with col3:
            st.metric("Warnings", total_warnings)
        
        # Display detailed results
        st.markdown("### 🔍 Detailed Validation Results")
        
        for result in validation_results:
            display_type = result['qtype'] if result['qtype'] else "MISSING TYPE"
            expander_title = f"❗ Question {result['qid']} (Row {result['row']}) - {display_type}"
            
            with st.expander(expander_title):
                
                # Show suggested type prominently
                if result.get('suggested_type'):
                    st.success(f"💡 **Suggested Type: {result['suggested_type']}** (Confidence: {result.get('confidence', 'N/A')})")
                
                if result['qtype'] and result['qtype'] in QTYPE_REQUIREMENTS:
                    st.info(f"ℹ️ **{result['qtype']}**: {QTYPE_REQUIREMENTS[result['qtype']]['description']}")
                
                if result['issues']:
                    st.markdown("**🔴 Issues:**")
                    for issue in result['issues']:
                        st.markdown(f"- {issue}")
                
                if result['warnings']:
                    st.markdown("**🟡 Warnings:**")
                    for warning in result['warnings']:
                        st.markdown(f"- {warning}")
                
                if result['requirements']:
                    st.markdown("**💡 Suggested Details Template:**")
                    if result['requirements'].get('details_template'):
                        st.code(result['requirements']['details_template'], language=None)
                    
                    st.markdown("**📋 Field Reference:**")
                    if result['requirements']['details_fields']['required']:
                        st.markdown(f"🔴 **Required:** {', '.join(result['requirements']['details_fields']['required'])}")
                    if result['requirements']['details_fields']['optional']:
                        st.markdown(f"🟡 **Optional:** {', '.join(result['requirements']['details_fields']['optional'])}")
        
        # Generate report document
        st.markdown("---")
        st.markdown("### 📄 Download Validation Report")
        st.markdown("Download your original document with a detailed validation report appended.")
        
        if st.button("Generate Report Document", type="primary"):
            with st.spinner("Generating report..."):
                # Create a new document from the original
                report_doc = Document(uploaded_file)
                
                # Add validation report
                report_doc = add_suggestions_to_document(report_doc, validation_results)
                
                # Save to bytes
                doc_bytes = io.BytesIO()
                report_doc.save(doc_bytes)
                doc_bytes.seek(0)
                
                st.download_button(
                    label="📥 Download Document with Validation Report",
                    data=doc_bytes,
                    file_name=f"validated_{uploaded_file.name}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("✅ Report generated! Original formatting preserved.")
